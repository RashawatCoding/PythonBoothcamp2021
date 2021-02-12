#test-docx.py
from docx import Document
import wikipedia

wikipedia.set_lang('th')

#summary สำหรับบทความที่่สรุป
data = wikipedia.summary('ประเทศไทย')

#page+content บทความทั้้งหน้า
data2 = wikipedia.page('ประเทศไทย')
data2 = data2.content

doc = Document() #สร้างไฟล์ เวิร์ด ใน python
doc.add_heading('แมว',0)

doc.add_paragraph(data2)

doc.save('ประเทศไทย.docx')
print('สร้างไฟล์สำเร็จ')
