#GUIWiki.py
import wikipedia

#python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    #summary สำหรับบทความที่่สรุป
    data = wikipedia.summary(keyword)

    #page+content บทความทั้้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์ เวิร์ด ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)

    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')

#config
FONT1 = ('Angsananew',15)

#คำอธิบาย
L = ttk.Label(GUI,text = 'ค้นหาบทความ',font = FONT1)
L.pack()

#ช่ิองค้นหาข้อมูล
v_search = StringVar() #กล่องสำหรับเก็บ key word
E1 = ttk.Entry(GUI,textvariable = v_search,font = FONT1,width = 35)
E1.pack(pady = 10)

#ปุ่มค้นหา
def search():
    keyword = v_search.get() #เป็นการดึงข้อมูลเข้ามา .get ใช้ได้กับ StraingVar() เท่านั้น
    try: #ลองค้นหาดูว่าได้ผลลัพท์หรือไม่ หากได้ให้ผ่านไป
        language = v_radio.get() # th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','บันทึกเรียบร้อยแล้ว')
    except: #หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')
        
    #print(wikipedia.search(keyword))#เป็นการบอกว่าคำที่เรา search มีทั้งหมดกี่คำ
    #result = wikipedia.summary(keyword)
    #print(result)
    
B1 = ttk.Button(GUI,text = 'Search',command = search)
B1.pack(ipadx = 20,ipady = 10)

#เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady = 10)

v_radio = StringVar() #ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text = 'ภาษาไทย',variable = v_radio,value = 'th')
RB2 = ttk.Radiobutton(F1,text = 'ภาษาอังกฤษ',variable = v_radio,value = 'en')
RB3 = ttk.Radiobutton(F1,text = 'ภาษาจีน',variable = v_radio,value = 'zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row = 0,column = 0)
RB2.grid(row = 0,column = 1)
RB3.grid(row = 0,column = 2)

GUI.mainloop()


